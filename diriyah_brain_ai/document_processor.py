"""
Document Processing Module for Diriyah Brain AI
Handles various document types: PDFs, CAD files, images, videos, KMZ files
"""
import os
import io
import json
import zipfile
import tempfile
from typing import Dict, List, Optional, Any, Union
from pathlib import Path
import logging
from datetime import datetime

# PDF Processing
try:
    import PyPDF2
    import pdfplumber
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Image Processing and OCR
try:
    from PIL import Image
    import pytesseract
    import cv2
    import numpy as np
    IMAGE_OCR_AVAILABLE = True
except ImportError:
    IMAGE_OCR_AVAILABLE = False

# CAD Processing
try:
    import ezdxf
    CAD_AVAILABLE = True
except ImportError:
    CAD_AVAILABLE = False

# Video Processing
try:
    import cv2
    VIDEO_AVAILABLE = True
except ImportError:
    VIDEO_AVAILABLE = False

# Excel/Spreadsheet Processing
try:
    import openpyxl
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# XML/KMZ Processing
try:
    import xml.etree.ElementTree as ET
    XML_AVAILABLE = True
except ImportError:
    XML_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Main document processing class"""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': ['pdf'],
            'image': ['jpg', 'jpeg', 'png', 'tiff', 'bmp', 'gif'],
            'cad': ['dxf', 'dwg'],
            'video': ['mp4', 'avi', 'mov', 'mkv', 'wmv'],
            'excel': ['xlsx', 'xls', 'csv'],
            'word': ['docx', 'doc'],
            'text': ['txt', 'md', 'rtf'],
            'kmz': ['kmz', 'kml'],
            'xml': ['xml'],
            'json': ['json']
        }
    
    def process_document(self, file_path: str, document_type: str = None) -> Dict[str, Any]:
        """
        Process a document and extract relevant information
        
        Args:
            file_path: Path to the document
            document_type: Type of document (optional, will be inferred if not provided)
            
        Returns:
            Dictionary containing extracted information
        """
        try:
            if not os.path.exists(file_path):
                return {'error': f'File not found: {file_path}'}
            
            file_extension = Path(file_path).suffix.lower().lstrip('.')
            
            if not document_type:
                document_type = self._infer_document_type(file_extension)
            
            logger.info(f"Processing {document_type} document: {file_path}")
            
            # Route to appropriate processor
            if document_type == 'pdf':
                return self._process_pdf(file_path)
            elif document_type == 'image':
                return self._process_image(file_path)
            elif document_type == 'cad':
                return self._process_cad(file_path)
            elif document_type == 'video':
                return self._process_video(file_path)
            elif document_type == 'excel':
                return self._process_excel(file_path)
            elif document_type == 'word':
                return self._process_word(file_path)
            elif document_type == 'text':
                return self._process_text(file_path)
            elif document_type == 'kmz':
                return self._process_kmz(file_path)
            elif document_type == 'xml':
                return self._process_xml(file_path)
            elif document_type == 'json':
                return self._process_json(file_path)
            else:
                return {'error': f'Unsupported document type: {document_type}'}
                
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            return {'error': f'Processing failed: {str(e)}'}
    
    def _infer_document_type(self, file_extension: str) -> str:
        """Infer document type from file extension"""
        for doc_type, extensions in self.supported_formats.items():
            if file_extension in extensions:
                return doc_type
        return 'unknown'
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF documents"""
        if not PDF_AVAILABLE:
            return {'error': 'PDF processing libraries not available'}
        
        result = {
            'type': 'pdf',
            'file_path': file_path,
            'text_content': '',
            'metadata': {},
            'pages': [],
            'tables': [],
            'images': [],
            'analysis': {}
        }
        
        try:
            # Extract text using pdfplumber (better for tables)
            with pdfplumber.open(file_path) as pdf:
                result['metadata'] = {
                    'pages': len(pdf.pages),
                    'title': pdf.metadata.get('Title', ''),
                    'author': pdf.metadata.get('Author', ''),
                    'subject': pdf.metadata.get('Subject', ''),
                    'creator': pdf.metadata.get('Creator', ''),
                    'creation_date': pdf.metadata.get('CreationDate', ''),
                }
                
                full_text = []
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ''
                    full_text.append(page_text)
                    
                    # Extract tables from each page
                    tables = page.extract_tables()
                    if tables:
                        for j, table in enumerate(tables):
                            result['tables'].append({
                                'page': i + 1,
                                'table_index': j + 1,
                                'data': table,
                                'rows': len(table),
                                'columns': len(table[0]) if table else 0
                            })
                    
                    result['pages'].append({
                        'page_number': i + 1,
                        'text': page_text,
                        'char_count': len(page_text),
                        'tables_count': len(tables) if tables else 0
                    })
                
                result['text_content'] = '\n'.join(full_text)
            
            # Additional processing with PyMuPDF for images and better metadata
            try:
                doc = fitz.open(file_path)
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    image_list = page.get_images()
                    for img_index, img in enumerate(image_list):
                        result['images'].append({
                            'page': page_num + 1,
                            'image_index': img_index + 1,
                            'xref': img[0],
                            'width': img[2],
                            'height': img[3]
                        })
                doc.close()
            except Exception as e:
                logger.warning(f"PyMuPDF processing failed: {e}")
            
            # Analyze content for construction-specific information
            result['analysis'] = self._analyze_construction_content(result['text_content'])
            
        except Exception as e:
            result['error'] = f'PDF processing failed: {str(e)}'
        
        return result
    
    def _process_image(self, file_path: str) -> Dict[str, Any]:
        """Process image files with OCR"""
        from diriyah_brain_ai.processors.media_processor import media_processor
        return media_processor.process_photo(file_path)

    def _process_video(self, file_path: str) -> Dict[str, Any]:
        """Process video files"""
        from diriyah_brain_ai.processors.media_processor import media_processor
        return media_processor.process_video(file_path)

    def _process_kmz(self, file_path: str) -> Dict[str, Any]:
        """Process KMZ/KML files"""
        from diriyah_brain_ai.processors.media_processor import media_processor
        return media_processor.process_kmz(file_path)

    
    def _process_cad(self, file_path: str) -> Dict[str, Any]:
        """Process CAD files (DXF format)"""
        if not CAD_AVAILABLE:
            return {'error': 'CAD processing libraries not available'}
        
        result = {
            'type': 'cad',
            'file_path': file_path,
            'metadata': {},
            'entities': [],
            'layers': [],
            'blocks': [],
            'analysis': {}
        }
        
        try:
            if file_path.lower().endswith('.dxf'):
                doc = ezdxf.readfile(file_path)
                
                result['metadata'] = {
                    'dxf_version': doc.dxfversion,
                    'units': doc.units,
                    'layers_count': len(doc.layers),
                    'blocks_count': len(doc.blocks),
                    'entities_count': len(list(doc.modelspace()))
                }
                
                # Extract layers
                for layer in doc.layers:
                    result['layers'].append({
                        'name': layer.dxf.name,
                        'color': layer.dxf.color,
                        'linetype': layer.dxf.linetype,
                        'on': not layer.is_off(),
                        'frozen': layer.is_frozen()
                    })
                
                # Extract entities from modelspace
                for entity in doc.modelspace():
                    entity_info = {
                        'type': entity.dxftype(),
                        'layer': entity.dxf.layer if hasattr(entity.dxf, 'layer') else 'Unknown'
                    }
                    
                    # Extract specific information based on entity type
                    if entity.dxftype() == 'TEXT':
                        entity_info['text'] = entity.dxf.text
                        entity_info['height'] = entity.dxf.height
                    elif entity.dxftype() == 'LINE':
                        entity_info['start'] = (entity.dxf.start.x, entity.dxf.start.y)
                        entity_info['end'] = (entity.dxf.end.x, entity.dxf.end.y)
                    elif entity.dxftype() == 'CIRCLE':
                        entity_info['center'] = (entity.dxf.center.x, entity.dxf.center.y)
                        entity_info['radius'] = entity.dxf.radius
                    
                    result['entities'].append(entity_info)
                
                # Extract text content for analysis
                text_entities = [e for e in result['entities'] if e.get('text')]
                text_content = '\n'.join([e['text'] for e in text_entities])
                result['analysis'] = self._analyze_construction_content(text_content)
                
            else:
                result['error'] = 'Only DXF files are currently supported for CAD processing'
                
        except Exception as e:
            result['error'] = f'CAD processing failed: {str(e)}'
        
        return result
    
    def _process_video(self, file_path: str) -> Dict[str, Any]:
        """Process video files"""
        if not VIDEO_AVAILABLE:
            return {'error': 'Video processing libraries not available'}
        
        result = {
            'type': 'video',
            'file_path': file_path,
            'metadata': {},
            'frames_analyzed': [],
            'analysis': {}
        }
        
        try:
            cap = cv2.VideoCapture(file_path)
            
            # Get video metadata
            result['metadata'] = {
                'fps': cap.get(cv2.CAP_PROP_FPS),
                'frame_count': int(cap.get(cv2.CAP_PROP_FRAME_COUNT)),
                'width': int(cap.get(cv2.CAP_PROP_FRAME_WIDTH)),
                'height': int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT)),
                'duration_seconds': int(cap.get(cv2.CAP_PROP_FRAME_COUNT)) / cap.get(cv2.CAP_PROP_FPS),
                'size_bytes': os.path.getsize(file_path)
            }
            
            # Sample frames for analysis (every 10% of the video)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            sample_frames = [int(frame_count * i / 10) for i in range(1, 10)]
            
            for frame_num in sample_frames:
                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_num)
                ret, frame = cap.read()
                
                if ret:
                    # Basic frame analysis
                    frame_info = {
                        'frame_number': frame_num,
                        'timestamp': frame_num / cap.get(cv2.CAP_PROP_FPS),
                        'brightness': np.mean(cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)),
                        'has_motion': True  # Placeholder for motion detection
                    }
                    result['frames_analyzed'].append(frame_info)
            
            cap.release()
            
            # Placeholder analysis
            result['analysis'] = {
                'video_type': 'construction_site',
                'quality': 'good' if result['metadata']['width'] >= 1280 else 'low',
                'suitable_for_progress_tracking': True
            }
            
        except Exception as e:
            result['error'] = f'Video processing failed: {str(e)}'
        
        return result
    
    def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """Process Excel/CSV files"""
        if not EXCEL_AVAILABLE:
            return {'error': 'Excel processing libraries not available'}
        
        result = {
            'type': 'excel',
            'file_path': file_path,
            'sheets': [],
            'metadata': {},
            'analysis': {}
        }
        
        try:
            if file_path.lower().endswith('.csv'):
                df = pd.read_csv(file_path)
                result['sheets'].append({
                    'name': 'Sheet1',
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': df.columns.tolist(),
                    'sample_data': df.head(5).to_dict('records')
                })
            else:
                # Excel file
                xl_file = pd.ExcelFile(file_path)
                result['metadata']['sheet_names'] = xl_file.sheet_names
                
                for sheet_name in xl_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    result['sheets'].append({
                        'name': sheet_name,
                        'rows': len(df),
                        'columns': len(df.columns),
                        'column_names': df.columns.tolist(),
                        'sample_data': df.head(5).to_dict('records')
                    })
            
            # Analyze for construction-specific content
            all_text = []
            for sheet in result['sheets']:
                all_text.extend([str(col) for col in sheet['column_names']])
                for row in sheet['sample_data']:
                    all_text.extend([str(val) for val in row.values() if val is not None])
            
            result['analysis'] = self._analyze_construction_content(' '.join(all_text))
            
        except Exception as e:
            result['error'] = f'Excel processing failed: {str(e)}'
        
        return result
    
    def _process_word(self, file_path: str) -> Dict[str, Any]:
        """Process Word documents"""
        try:
            from docx import Document
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            text_content = "\n".join(full_text)
            
            result = {
                "type": "word",
                "file_path": file_path,
                "text_content": text_content,
                "metadata": {
                    "paragraphs": len(doc.paragraphs),
                    "words": len(text_content.split())
                },
                "analysis": self._analyze_construction_content(text_content)
            }
        except ImportError:
            result = {"error": "python-docx library not available for Word processing"}
        except Exception as e:
            result = {"error": f"Word processing failed: {str(e)}"}
        return result
    
    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process plain text files"""
        result = {
            'type': 'text',
            'file_path': file_path,
            'text_content': '',
            'metadata': {},
            'analysis': {}
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            result['text_content'] = content
            result['metadata'] = {
                'size_bytes': os.path.getsize(file_path),
                'line_count': len(content.split('\n')),
                'word_count': len(content.split()),
                'char_count': len(content)
            }
            
            result['analysis'] = self._analyze_construction_content(content)
            
        except Exception as e:
            result['error'] = f'Text processing failed: {str(e)}'
        
        return result
    
    def _process_kmz(self, file_path: str) -> Dict[str, Any]:
        """Process KMZ/KML files"""
        result = {
            'type': 'kmz',
            'file_path': file_path,
            'placemarks': [],
            'metadata': {},
            'analysis': {}
        }
        
        try:
            if file_path.lower().endswith('.kmz'):
                # Extract KML from KMZ
                with zipfile.ZipFile(file_path, 'r') as kmz:
                    kml_files = [f for f in kmz.namelist() if f.endswith('.kml')]
                    if kml_files:
                        kml_content = kmz.read(kml_files[0]).decode('utf-8')
                    else:
                        return {'error': 'No KML file found in KMZ'}
            else:
                # Direct KML file
                with open(file_path, 'r', encoding='utf-8') as f:
                    kml_content = f.read()
            
            # Parse KML
            root = ET.fromstring(kml_content)
            
            # Extract placemarks
            for placemark in root.iter():
                if placemark.tag.endswith('Placemark'):
                    name = placemark.find('.//{http://www.opengis.net/kml/2.2}name')
                    description = placemark.find('.//{http://www.opengis.net/kml/2.2}description')
                    coordinates = placemark.find('.//{http://www.opengis.net/kml/2.2}coordinates')
                    
                    result['placemarks'].append({
                        'name': name.text if name is not None else 'Unnamed',
                        'description': description.text if description is not None else '',
                        'coordinates': coordinates.text if coordinates is not None else ''
                    })
            
            result['metadata']['placemarks_count'] = len(result['placemarks'])
            
            # Analyze descriptions for construction content
            all_text = ' '.join([p['name'] + ' ' + p['description'] for p in result['placemarks']])
            result['analysis'] = self._analyze_construction_content(all_text)
            
        except Exception as e:
            result['error'] = f'KMZ/KML processing failed: {str(e)}'
        
        return result
    
    def _process_xml(self, file_path: str) -> Dict[str, Any]:
        """Process XML files"""
        result = {
            'type': 'xml',
            'file_path': file_path,
            'structure': {},
            'text_content': '',
            'analysis': {}
        }
        
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            result['structure'] = {
                'root_tag': root.tag,
                'attributes': root.attrib,
                'children_count': len(root)
            }
            
            # Extract all text content
            all_text = []
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    all_text.append(elem.text.strip())
            
            result['text_content'] = '\n'.join(all_text)
            result['analysis'] = self._analyze_construction_content(result['text_content'])
            
        except Exception as e:
            result['error'] = f'XML processing failed: {str(e)}'
        
        return result
    
    def _process_json(self, file_path: str) -> Dict[str, Any]:
        """Process JSON files"""
        result = {
            'type': 'json',
            'file_path': file_path,
            'data': {},
            'text_content': '',
            'analysis': {}
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            result['data'] = data
            
            # Extract text content for analysis
            def extract_text(obj):
                texts = []
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        texts.append(str(key))
                        texts.extend(extract_text(value))
                elif isinstance(obj, list):
                    for item in obj:
                        texts.extend(extract_text(item))
                elif isinstance(obj, str):
                    texts.append(obj)
                else:
                    texts.append(str(obj))
                return texts
            
            all_text = extract_text(data)
            result['text_content'] = ' '.join(all_text)
            result['analysis'] = self._analyze_construction_content(result['text_content'])
            
        except Exception as e:
            result['error'] = f'JSON processing failed: {str(e)}'
        
        return result
    
    def _analyze_construction_content(self, text: str) -> Dict[str, Any]:
        """Analyze text content for construction-specific information"""
        if not text:
            return {}
        
        text_lower = text.lower()
        
        # Construction-specific keywords and patterns
        keywords = {
            'project_phases': ['foundation', 'structure', 'finishing', 'mep', 'landscaping', 'handover'],
            'materials': ['concrete', 'steel', 'rebar', 'asphalt', 'aggregate', 'cement', 'sand'],
            'equipment': ['crane', 'excavator', 'bulldozer', 'mixer', 'pump', 'compactor'],
            'documents': ['boq', 'schedule', 'rfi', 'ncr', 'mom', 'variation', 'contract'],
            'quality': ['inspection', 'test', 'compliance', 'standard', 'specification'],
            'safety': ['safety', 'hazard', 'ppe', 'incident', 'risk', 'hse'],
            'financial': ['cost', 'budget', 'payment', 'invoice', 'claim', 'variation']
        }
        
        analysis = {
            'document_category': 'unknown',
            'keywords_found': {},
            'likely_content_type': [],
            'language': 'mixed' if any('\u0600' <= char <= '\u06FF' for char in text) else 'english',
            'confidence_score': 0.0
        }
        
        # Count keyword occurrences
        total_matches = 0
        for category, words in keywords.items():
            matches = sum(1 for word in words if word in text_lower)
            if matches > 0:
                analysis['keywords_found'][category] = matches
                total_matches += matches
        
        # Determine likely content type
        if analysis['keywords_found']:
            max_category = max(analysis['keywords_found'], key=analysis['keywords_found'].get)
            analysis['likely_content_type'].append(max_category)
            
            # Specific document type detection
            if 'boq' in text_lower or 'bill of quantities' in text_lower:
                analysis['document_category'] = 'boq'
            elif 'schedule' in text_lower and ('gantt' in text_lower or 'milestone' in text_lower):
                analysis['document_category'] = 'schedule'
            elif 'rfi' in text_lower or 'request for information' in text_lower:
                analysis['document_category'] = 'rfi'
            elif 'ncr' in text_lower or 'non-conformance' in text_lower:
                analysis['document_category'] = 'ncr'
            elif 'contract' in text_lower and ('agreement' in text_lower or 'terms' in text_lower):
                analysis['document_category'] = 'contract'
            elif 'mom' in text_lower or 'minutes of meeting' in text_lower:
                analysis['document_category'] = 'mom'
        
        # Calculate confidence score
        analysis['confidence_score'] = min(total_matches / 10.0, 1.0)  # Normalize to 0-1
        
        return analysis

# Global instance
document_processor = DocumentProcessor()


         # Analyze for construction-specific content
            text_content = "\n".join([str(item) for sheet in result["sheets"] for row in sheet["sample_data"] for item in row.values()])
            result["analysis"] = self._analyze_construction_content(text_content)
            
        except Exception as e:
            result["error"] = f"Excel processing failed: {str(e)}"
        
        return result
    
    def _process_word(self, file_path: str) -> Dict[str, Any]:
        """Process Word documents"""
        try:
            from docx import Document
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            text_content = "\n".join(full_text)
            
            result = {
                "type": "word",
                "file_path": file_path,
                "text_content": text_content,
                "metadata": {
                    "paragraphs": len(doc.paragraphs),
                    "words": len(text_content.split())
                },
                "analysis": self._analyze_construction_content(text_content)
            }
        except ImportError:
            result = {"error": "python-docx library not available for Word processing"}
        except Exception as e:
            result = {"error": f"Word processing failed: {str(e)}"}
        return result
    
    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process plain text files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text_content = f.read()
            
            result = {
                "type": "text",
                "file_path": file_path,
                "text_content": text_content,
                "metadata": {
                    "char_count": len(text_content),
                    "word_count": len(text_content.split())
                },
                "analysis": self._analyze_construction_content(text_content)
            }
        except Exception as e:
            result = {"error": f"Text file processing failed: {str(e)}"}
        return result
    
    def _process_kmz(self, file_path: str) -> Dict[str, Any]:
        """Process KMZ/KML files"""
        if not XML_AVAILABLE:
            return {"error": "XML processing libraries not available"}
        
        result = {
            "type": "kmz",
            "file_path": file_path,
            "kml_content": "",
            "placemarks": [],
            "metadata": {},
            "analysis": {}
        }
        
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                with zipfile.ZipFile(file_path, "r") as zip_ref:
                    zip_ref.extractall(tmpdir)
                
                kml_file = None
                for root, _, files in os.walk(tmpdir):
                    for file in files:
                        if file.endswith(".kml"):
                            kml_file = os.path.join(root, file)
                            break
                    if kml_file: break
                
                if not kml_file:
                    raise ValueError("KML file not found inside KMZ archive")
                
                with open(kml_file, "r", encoding="utf-8") as f:
                    kml_content = f.read()
                result["kml_content"] = kml_content
                
                # Parse KML for placemarks
                root = ET.fromstring(kml_content)
                for placemark in root.findall(".//{http://www.opengis.net/kml/2.2}Placemark"):
                    name = placemark.find(".//{http://www.opengis.net/kml/2.2}name")
                    description = placemark.find(".//{http://www.opengis.net/kml/2.2}description")
                    point = placemark.find(".//{http://www.opengis.net/kml/2.2}Point/{http://www.opengis.net/kml/2.2}coordinates")
                    
                    placemark_data = {
                        "name": name.text if name is not None else "",
                        "description": description.text if description is not None else "",
                        "coordinates": point.text.strip() if point is not None else ""
                    }
                    result["placemarks"].append(placemark_data)
            
            result["metadata"]["placemark_count"] = len(result["placemarks"])
            result["analysis"] = self._analyze_construction_content(kml_content)
            
        except Exception as e:
            result["error"] = f"KMZ processing failed: {str(e)}"
        
        return result
    
    def _process_xml(self, file_path: str) -> Dict[str, Any]:
        """Process XML files"""
        if not XML_AVAILABLE:
            return {"error": "XML processing libraries not available"}
        
        result = {
            "type": "xml",
            "file_path": file_path,
            "root_element": "",
            "elements_count": 0,
            "text_content": "",
            "analysis": {}
        }
        
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            result["root_element"] = root.tag
            
            # Extract all text content
            text_content = ". ".join([elem.text for elem in root.iter() if elem.text and elem.text.strip()])
            result["text_content"] = text_content
            result["elements_count"] = len(list(root.iter()))
            result["analysis"] = self._analyze_construction_content(text_content)
            
        except Exception as e:
            result["error"] = f"XML processing failed: {str(e)}"
        
        return result
    
    def _process_json(self, file_path: str) -> Dict[str, Any]:
        """Process JSON files"""
        result = {
            "type": "json",
            "file_path": file_path,
            "data": {},
            "analysis": {}
        }
        
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            result["data"] = data
            result["analysis"] = self._analyze_construction_content(json.dumps(data))
            
        except Exception as e:
            result["error"] = f"JSON processing failed: {str(e)}"
        
        return result
    
    def _analyze_construction_content(self, text_content: str) -> Dict[str, Any]:
        """Analyze text content for construction-specific keywords and entities"""
        analysis = {
            "keywords": [],
            "entities": [],
            "sentiment": "neutral",  # Placeholder
            "language": "english",  # Placeholder
            "document_category": "general"
        }
        
        text_upper = text_content.upper()
        
        # Simple keyword extraction
        construction_keywords = [
            "BOQ", "SCHEDULE", "RFI", "NCR", "CONTRACT", "DRAWING", "SPECIFICATION",
            "CONCRETE", "STEEL", "EXCAVATION", "FOUNDATION", "MEP", "FINISHES",
            "SAFETY", "QUALITY", "RISK", "DELAY", "COST", "BUDGET", "VARIATION"
        ]
        
        found_keywords = [kw for kw in construction_keywords if kw in text_upper]
        analysis["keywords"] = list(set(found_keywords))
        
        # Document categorization (simplified)
        if "BOQ" in text_upper or "BILL OF QUANTITIES" in text_upper:
            analysis["document_category"] = "BOQ"
        elif "SCHEDULE" in text_upper or "PROGRAMME" in text_upper or "P6" in text_upper:
            analysis["document_category"] = "Schedule"
        elif "RFI" in text_upper or "REQUEST FOR INFORMATION" in text_upper:
            analysis["document_category"] = "RFI"
        elif "NCR" in text_upper or "NON-CONFORMANCE REPORT" in text_upper:
            analysis["document_category"] = "NCR"
        elif "CONTRACT" in text_upper or "AGREEMENT" in text_upper:
            analysis["document_category"] = "Contract"
        elif "SPECIFICATION" in text_upper or "TECHNICAL REQUIREMENTS" in text_upper:
            analysis["document_category"] = "Specification"
        elif "DRAWING" in text_upper or "PLAN" in text_upper or "CAD" in text_upper:
            analysis["document_category"] = "Drawing"
        
        # Placeholder for more advanced NLP for entities and sentiment
        
        return analysis

# Global instance
document_processor = DocumentProcessor()

# Import specific processors
from diriyah_brain_ai.processors.cad_processor import cad_processor
from diriyah_brain_ai.processors.p6_processor import p6_processor
from diriyah_brain_ai.processors.aconex_processor import aconex_processor

class AdvancedDocumentProcessor(DocumentProcessor):
    def _process_cad(self, file_path: str) -> Dict[str, Any]:
        return cad_processor.process_cad_file(file_path)

    def _process_p6(self, file_path: str) -> Dict[str, Any]:
        return p6_processor.process_p6_file(file_path)

    def _process_aconex(self, file_path: str) -> Dict[str, Any]:
        # Aconex extracts are typically text files or PDFs, so we'll read text and then process
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text_content = f.read()
            return aconex_processor.process_aconex_extract(text_content, os.path.basename(file_path))
        except Exception as e:
            return {"error": f"Aconex processing failed: {str(e)}"}

    def _process_bim(self, file_path: str) -> Dict[str, Any]:
        from diriyah_brain_ai.processors.bim_processor import bim_processor
        return bim_processor.process_bim_file(file_path)

    def _process_powerbi(self, file_path: str) -> Dict[str, Any]:
        from diriyah_brain_ai.processors.powerbi_processor import powerbi_processor
        return powerbi_processor.process_powerbi_report(file_path)

# Replace the global instance with the advanced one
document_processor = AdvancedDocumentProcessor()

