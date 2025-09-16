"""
Enhanced Google Drive Client for Diriyah Brain AI
Comprehensive file access and analysis including ZIP/RAR archives
"""

import os
import io
import zipfile
import rarfile
import tempfile
import mimetypes
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import logging
from datetime import datetime

from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import PyPDF2
import docx
from PIL import Image
import pandas as pd

logger = logging.getLogger(__name__)

class EnhancedDriveClient:
    """Enhanced Google Drive client with comprehensive file analysis capabilities"""
    
    def __init__(self):
        self.service = None
        self.credentials = None
        self.supported_formats = {
            # Documents
            'application/pdf': self._analyze_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._analyze_docx,
            'application/msword': self._analyze_doc,
            'text/plain': self._analyze_text,
            'text/csv': self._analyze_csv,
            
            # Spreadsheets
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._analyze_excel,
            'application/vnd.ms-excel': self._analyze_excel,
            
            # Images
            'image/jpeg': self._analyze_image,
            'image/png': self._analyze_image,
            'image/gif': self._analyze_image,
            'image/bmp': self._analyze_image,
            
            # CAD Files
            'application/acad': self._analyze_cad,
            'application/x-autocad': self._analyze_cad,
            'image/vnd.dwg': self._analyze_cad,
            
            # Archives
            'application/zip': self._analyze_zip,
            'application/x-rar-compressed': self._analyze_rar,
            'application/x-rar': self._analyze_rar,
            
            # Google Workspace
            'application/vnd.google-apps.document': self._analyze_google_doc,
            'application/vnd.google-apps.spreadsheet': self._analyze_google_sheet,
            'application/vnd.google-apps.presentation': self._analyze_google_slides,
        }
    
    def authenticate(self, client_id: str, client_secret: str, redirect_uri: str) -> str:
        """Initialize OAuth flow and return authorization URL"""
        try:
            flow = Flow.from_client_config(
                {
                    "web": {
                        "client_id": client_id,
                        "client_secret": client_secret,
                        "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                        "token_uri": "https://oauth2.googleapis.com/token",
                        "redirect_uris": [redirect_uri]
                    }
                },
                scopes=['https://www.googleapis.com/auth/drive.readonly']
            )
            flow.redirect_uri = redirect_uri
            
            auth_url, _ = flow.authorization_url(prompt='consent')
            return auth_url
            
        except Exception as e:
            logger.error(f"Authentication error: {e}")
            raise
    
    def complete_auth(self, authorization_code: str, client_id: str, client_secret: str, redirect_uri: str):
        """Complete OAuth flow with authorization code"""
        try:
            flow = Flow.from_client_config(
                {
                    "web": {
                        "client_id": client_id,
                        "client_secret": client_secret,
                        "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                        "token_uri": "https://oauth2.googleapis.com/token",
                        "redirect_uris": [redirect_uri]
                    }
                },
                scopes=['https://www.googleapis.com/auth/drive.readonly']
            )
            flow.redirect_uri = redirect_uri
            
            flow.fetch_token(code=authorization_code)
            self.credentials = flow.credentials
            self.service = build('drive', 'v3', credentials=self.credentials)
            
            logger.info("Google Drive authentication completed successfully")
            
        except Exception as e:
            logger.error(f"Auth completion error: {e}")
            raise
    
    def scan_all_files(self, include_shared: bool = True) -> List[Dict[str, Any]]:
        """Comprehensively scan all accessible files and folders"""
        try:
            all_files = []
            page_token = None
            
            while True:
                # Build query to include shared files
                query = "trashed=false"
                if include_shared:
                    query += " or sharedWithMe=true"
                
                results = self.service.files().list(
                    q=query,
                    pageSize=1000,
                    fields="nextPageToken, files(id, name, mimeType, size, createdTime, modifiedTime, parents, shared, owners, permissions, webViewLink)",
                    pageToken=page_token,
                    includeItemsFromAllDrives=True,
                    supportsAllDrives=True
                ).execute()
                
                files = results.get('files', [])
                all_files.extend(files)
                
                page_token = results.get('nextPageToken')
                if not page_token:
                    break
            
            logger.info(f"Scanned {len(all_files)} files from Google Drive")
            return all_files
            
        except Exception as e:
            logger.error(f"File scanning error: {e}")
            return []
    
    def analyze_file(self, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze a single file comprehensively"""
        try:
            file_id = file_info['id']
            file_name = file_info['name']
            mime_type = file_info['mimeType']
            
            analysis = {
                'file_id': file_id,
                'name': file_name,
                'mime_type': mime_type,
                'size': file_info.get('size', 0),
                'created': file_info.get('createdTime'),
                'modified': file_info.get('modifiedTime'),
                'shared': file_info.get('shared', False),
                'web_link': file_info.get('webViewLink'),
                'analysis_status': 'pending',
                'content_summary': '',
                'extracted_text': '',
                'metadata': {},
                'errors': []
            }
            
            # Skip folders for content analysis
            if mime_type == 'application/vnd.google-apps.folder':
                analysis['analysis_status'] = 'folder'
                analysis['content_summary'] = 'Folder container'
                return analysis
            
            # Analyze based on file type
            if mime_type in self.supported_formats:
                try:
                    content_analysis = self.supported_formats[mime_type](file_id, file_name)
                    analysis.update(content_analysis)
                    analysis['analysis_status'] = 'completed'
                except Exception as e:
                    analysis['errors'].append(f"Analysis error: {str(e)}")
                    analysis['analysis_status'] = 'error'
            else:
                analysis['analysis_status'] = 'unsupported'
                analysis['content_summary'] = f'Unsupported file type: {mime_type}'
            
            return analysis
            
        except Exception as e:
            logger.error(f"File analysis error for {file_info.get('name', 'unknown')}: {e}")
            return {
                'file_id': file_info.get('id'),
                'name': file_info.get('name', 'unknown'),
                'analysis_status': 'error',
                'errors': [str(e)]
            }
    
    def download_file(self, file_id: str) -> bytes:
        """Download file content as bytes"""
        try:
            request = self.service.files().get_media(fileId=file_id)
            file_io = io.BytesIO()
            downloader = MediaIoBaseDownload(file_io, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            
            return file_io.getvalue()
            
        except Exception as e:
            logger.error(f"Download error for file {file_id}: {e}")
            raise
    
    def _analyze_pdf(self, file_id: str, file_name: str) -> Dict[str, Any]:
        """Analyze PDF files"""
        try:
            content = self.download_file(file_id)
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            
            return {
                'extracted_text': text_content[:5000],  # Limit for storage
                'content_summary': f'PDF document with {len(pdf_reader.pages)} pages',
                'metadata': {
                    'pages': len(pdf_reader.pages),
                    'file_type': 'PDF Document'
                }
            }
            
        except Exception as e:
            return {
                'extracted_text': '',
                'content_summary': 'PDF analysis failed',
                'errors': [str(e)]
            }
    
    def _analyze_docx(self, file_id: str, file_name: str) -> Dict[str, Any]:
        """Analyze Word documents"""
        try:
            content = self.download_file(file_id)
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            return {
                'extracted_text': text_content[:5000],
                'content_summary': f'Word document with {len(doc.paragraphs)} paragraphs',
                'metadata': {
                    'paragraphs': len(doc.paragraphs),
                    'file_type': 'Word Document'
                }
            }
            
        except Exception as e:
            return {
                'extracted_text': '',
                'content_summary': 'Word document analysis failed',
                'errors': [str(e)]
            }
    
    def _analyze_text(self, file_id: str, file_name: str) -> Dict[str, Any]:
        """Analyze plain text files"""
        try:
            content = self.download_file(file_id)
            text_content = content.decode('utf-8', errors='ignore')
            
            lines = text_content.split('\n')
            
            return {
                'extracted_text': text_content[:5000],
                'content_summary': f'Text file with {len(lines)} lines',
                'metadata': {
                    'lines': len(lines),
                    'characters': len(text_content),
                    'file_type': 'Text File'
                }
            }
            
        except Exception as e:
            return {
                'extracted_text': '',
                'content_summary': 'Text file analysis failed',
                'errors': [str(e)]
            }
    
    def _analyze_csv(self, file_id: str, file_name: str) -> Dict[str, Any]:
        """Analyze CSV files"""
        try:
            content = self.download_file(file_id)
            csv_file = io.StringIO(content.decode('utf-8', errors='ignore'))
            df = pd.read_csv(csv_file)
            
            return {
                'extracted_text': df.head(10).to_string(),
                'content_summary': f'CSV file with {len(df)} rows and {len(df.columns)} columns',
                'metadata': {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': list(df.columns),
                    'file_type': 'CSV Data'
                }
            }
            
        except Exception as e:
            return {
                'extracted_text': '',
                'content_summary': 'CSV analysis failed',
                'errors': [str(e)]
            }
    
    def _analyze_excel(self, file_id: str, file_name: str) -> Dict[str, Any]:
        """Analyze Excel files"""
        try:
            content = self.download_file(file_id)
            excel_file = io.BytesIO(content)
            df = pd.read_excel(excel_file, sheet_name=None)
            
            sheets_info = {}
            total_rows = 0
            
            for sheet_name, sheet_df in df.items():
                sheets_info[sheet_name] = {
                    'rows': len(sheet_df),
                    'columns': len(sheet_df.columns)
                }
                total_rows += len(sheet_df)
            
            # Get sample data from first sheet
            first_sheet = list(df.values())[0] if df else pd.DataFrame()
            sample_text = first_sheet.head(5).to_string() if not first_sheet.empty else ""
            
            return {
                'extracted_text': sample_text,
                'content_summary': f'Excel file with {len(df)} sheets and {total_rows} total rows',
                'metadata': {
                    'sheets': len(df),
                    'total_rows': total_rows,
                    'sheets_info': sheets_info,
                    'file_type': 'Excel Spreadsheet'
                }
            }
            
        except Exception as e:
            return {
                'extracted_text': '',
                'content_summary': 'Excel analysis failed',
                'errors': [str(e)]
            }
    
    def _analyze_image(self, file_id: str, file_name: str) -> Dict[str, Any]:
        """Analyze image files"""
        try:
            content = self.download_file(file_id)
            image_file = io.BytesIO(content)
            image = Image.open(image_file)
            
            return {
                'extracted_text': f'Image file: {file_name}',
                'content_summary': f'Image file ({image.format}) - {image.size[0]}x{image.size[1]} pixels',
                'metadata': {
                    'format': image.format,
                    'size': image.size,
                    'mode': image.mode,
                    'file_type': 'Image File'
                }
            }
            
        except Exception as e:
            return {
                'extracted_text': '',
                'content_summary': 'Image analysis failed',
                'errors': [str(e)]
            }
    
    def _analyze_cad(self, file_id: str, file_name: str) -> Dict[str, Any]:
        """Analyze CAD files (basic metadata)"""
        try:
            content = self.download_file(file_id)
            
            return {
                'extracted_text': f'CAD file: {file_name}',
                'content_summary': f'CAD drawing file - {len(content)} bytes',
                'metadata': {
                    'file_size': len(content),
                    'file_type': 'CAD Drawing'
                }
            }
            
        except Exception as e:
            return {
                'extracted_text': '',
                'content_summary': 'CAD file analysis failed',
                'errors': [str(e)]
            }
    
    def _analyze_zip(self, file_id: str, file_name: str) -> Dict[str, Any]:
        """Analyze ZIP archives"""
        try:
            content = self.download_file(file_id)
            zip_file = io.BytesIO(content)
            
            with zipfile.ZipFile(zip_file, 'r') as zip_ref:
                file_list = zip_ref.namelist()
                
                # Analyze contents
                extracted_content = ""
                file_types = {}
                
                for file_path in file_list[:20]:  # Limit to first 20 files
                    try:
                        if not file_path.endswith('/'):  # Skip directories
                            file_ext = Path(file_path).suffix.lower()
                            file_types[file_ext] = file_types.get(file_ext, 0) + 1
                            
                            # Try to extract text from text files
                            if file_ext in ['.txt', '.md', '.py', '.js', '.html', '.css']:
                                with zip_ref.open(file_path) as f:
                                    try:
                                        text = f.read().decode('utf-8', errors='ignore')[:1000]
                                        extracted_content += f"\n--- {file_path} ---\n{text}\n"
                                    except:
                                        pass
                    except:
                        continue
                
                return {
                    'extracted_text': extracted_content[:5000],
                    'content_summary': f'ZIP archive with {len(file_list)} files',
                    'metadata': {
                        'total_files': len(file_list),
                        'file_types': file_types,
                        'file_list': file_list[:50],  # First 50 files
                        'file_type': 'ZIP Archive'
                    }
                }
                
        except Exception as e:
            return {
                'extracted_text': '',
                'content_summary': 'ZIP analysis failed',
                'errors': [str(e)]
            }
    
    def _analyze_rar(self, file_id: str, file_name: str) -> Dict[str, Any]:
        """Analyze RAR archives"""
        try:
            content = self.download_file(file_id)
            
            # Save to temporary file (rarfile requires file path)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.rar') as temp_file:
                temp_file.write(content)
                temp_path = temp_file.name
            
            try:
                with rarfile.RarFile(temp_path) as rar_ref:
                    file_list = rar_ref.namelist()
                    
                    # Analyze contents
                    file_types = {}
                    for file_path in file_list:
                        if not file_path.endswith('/'):
                            file_ext = Path(file_path).suffix.lower()
                            file_types[file_ext] = file_types.get(file_ext, 0) + 1
                    
                    return {
                        'extracted_text': f'RAR archive contents: {", ".join(file_list[:10])}',
                        'content_summary': f'RAR archive with {len(file_list)} files',
                        'metadata': {
                            'total_files': len(file_list),
                            'file_types': file_types,
                            'file_list': file_list[:50],
                            'file_type': 'RAR Archive'
                        }
                    }
            finally:
                os.unlink(temp_path)
                
        except Exception as e:
            return {
                'extracted_text': '',
                'content_summary': 'RAR analysis failed',
                'errors': [str(e)]
            }
    
    def _analyze_google_doc(self, file_id: str, file_name: str) -> Dict[str, Any]:
        """Analyze Google Docs"""
        try:
            # Export as plain text
            request = self.service.files().export_media(fileId=file_id, mimeType='text/plain')
            file_io = io.BytesIO()
            downloader = MediaIoBaseDownload(file_io, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            
            text_content = file_io.getvalue().decode('utf-8', errors='ignore')
            
            return {
                'extracted_text': text_content[:5000],
                'content_summary': f'Google Doc with {len(text_content.split())} words',
                'metadata': {
                    'words': len(text_content.split()),
                    'characters': len(text_content),
                    'file_type': 'Google Document'
                }
            }
            
        except Exception as e:
            return {
                'extracted_text': '',
                'content_summary': 'Google Doc analysis failed',
                'errors': [str(e)]
            }
    
    def _analyze_google_sheet(self, file_id: str, file_name: str) -> Dict[str, Any]:
        """Analyze Google Sheets"""
        try:
            # Export as CSV
            request = self.service.files().export_media(fileId=file_id, mimeType='text/csv')
            file_io = io.BytesIO()
            downloader = MediaIoBaseDownload(file_io, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            
            csv_content = file_io.getvalue().decode('utf-8', errors='ignore')
            df = pd.read_csv(io.StringIO(csv_content))
            
            return {
                'extracted_text': df.head(10).to_string(),
                'content_summary': f'Google Sheet with {len(df)} rows and {len(df.columns)} columns',
                'metadata': {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': list(df.columns),
                    'file_type': 'Google Spreadsheet'
                }
            }
            
        except Exception as e:
            return {
                'extracted_text': '',
                'content_summary': 'Google Sheet analysis failed',
                'errors': [str(e)]
            }
    
    def _analyze_google_slides(self, file_id: str, file_name: str) -> Dict[str, Any]:
        """Analyze Google Slides"""
        try:
            # Export as plain text
            request = self.service.files().export_media(fileId=file_id, mimeType='text/plain')
            file_io = io.BytesIO()
            downloader = MediaIoBaseDownload(file_io, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            
            text_content = file_io.getvalue().decode('utf-8', errors='ignore')
            
            return {
                'extracted_text': text_content[:5000],
                'content_summary': f'Google Slides presentation',
                'metadata': {
                    'content_length': len(text_content),
                    'file_type': 'Google Presentation'
                }
            }
            
        except Exception as e:
            return {
                'extracted_text': '',
                'content_summary': 'Google Slides analysis failed',
                'errors': [str(e)]
            }
    
    def _analyze_doc(self, file_id: str, file_name: str) -> Dict[str, Any]:
        """Analyze legacy Word documents"""
        return {
            'extracted_text': f'Legacy Word document: {file_name}',
            'content_summary': 'Legacy Word document (requires specialized processing)',
            'metadata': {
                'file_type': 'Legacy Word Document'
            }
        }
    
    def get_folder_structure(self) -> Dict[str, Any]:
        """Get complete folder structure"""
        try:
            folders = self.service.files().list(
                q="mimeType='application/vnd.google-apps.folder' and trashed=false",
                fields="files(id, name, parents)",
                pageSize=1000
            ).execute().get('files', [])
            
            # Build folder tree
            folder_tree = {}
            for folder in folders:
                folder_tree[folder['id']] = {
                    'name': folder['name'],
                    'parents': folder.get('parents', []),
                    'children': []
                }
            
            return folder_tree
            
        except Exception as e:
            logger.error(f"Folder structure error: {e}")
            return {}
    
    def search_content(self, query: str, file_analyses: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Search through analyzed content"""
        results = []
        query_lower = query.lower()
        
        for analysis in file_analyses:
            score = 0
            
            # Search in file name
            if query_lower in analysis.get('name', '').lower():
                score += 10
            
            # Search in extracted text
            extracted_text = analysis.get('extracted_text', '').lower()
            if query_lower in extracted_text:
                score += 5
                
            # Search in content summary
            if query_lower in analysis.get('content_summary', '').lower():
                score += 3
            
            if score > 0:
                analysis['search_score'] = score
                results.append(analysis)
        
        # Sort by relevance
        results.sort(key=lambda x: x['search_score'], reverse=True)
        return results

